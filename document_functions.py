from docx import Document
from docx.shared import Inches
import json
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage: set_cell_border(
        cell,
        top={"sz": 0, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 0, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 0, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    if not hasattr(tcPr, 'tcBorders'):
        tcPr.tcBorders = OxmlElement('w:tcBorders')

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            if not hasattr(tcPr.tcBorders, tag):
                el = OxmlElement(tag)
                tcPr.tcBorders.append(el)
            else:
                el = tcPr.tcBorders.__getattr__(tag)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    el.set(key, str(edge_data[key]))

def create_word_document_from_json(json_data, title="Course_Outline.docx"):

    doc = Document()
    doc.add_heading('Course Outline', level=1)

    # Create a table for course details
    details_table = doc.add_table(rows=2, cols=2)

    # Set table style to enable borders
    details_table.style = 'Table Grid'

    # Add course details in two columns
    detail_keys = ["course_title", "instructor_name", "credit_units", "total_hours"]
    for i, key in enumerate(detail_keys):
        cell = details_table.cell(i // 2, i % 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.text = f"{key.replace('_', ' ').capitalize()}: {json_data[key]}"
        # Make the detail name bold
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True

    # Add Course Description
    doc.add_heading("Course Description:", level=2)
    doc.add_paragraph(json_data["course_description"], style='BodyText')

    # Add Course Learning Outcomes (CLOs)
    doc.add_heading('Course Learning Outcomes (CLOs)', level=2)
    doc.add_paragraph("By the end of this course, students will be able to:")
    for i, clo in enumerate(json_data['clos'], 1):
        doc.add_paragraph(f"CLO {i}: {clo}", style='ListBullet')

    # Add Topics / Modules and Intended Learning Outcomes (ILOs)
    doc.add_heading('Topics / Modules and Intended Learning Outcomes', level=2)
    for i,topic in enumerate(json_data['topics']):
        doc.add_paragraph(f"Topic {i+1}: {topic['topic']}", style='ListNumber')
        for j,ilo in enumerate(topic['ilos']):
            doc.add_paragraph(f"ILO {i+1}.{j+1}: {ilo}", style='ListBullet2')
    
    # Add a table to the document
    doc.add_heading('Weekly Activities', level=2)
    table = doc.add_table(rows=1, cols=5)

    # Set table style to enable borders
    table.style = 'Table Grid'

    # Add headers to the table
    hdr_cells = table.rows[0].cells
    headers = ["Week No.", "Topic", "Activity Description", "Expected Output", "Assessment Tools"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs
        run[0].font.bold = True
        run[0].font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths
        if header == "Week No.":
            hdr_cells[i].width = Inches(0.75)
        elif header == "Activity Description":
            hdr_cells[i].width = Inches(3)
        else:
            hdr_cells[i].width = Inches(1.5)

    # Add rows to the table from the JSON data
    for activity in json_data['activities']:
        row_cells = table.add_row().cells
        row_cells[0].text = activity['week']
        row_cells[1].text = activity['topic']
        row_cells[2].text = activity['activity_description']
        row_cells[3].text = activity['expected_output']
        row_cells[4].text = activity['assessment_tools']

        for cell in row_cells:
            set_cell_border(cell, top={"sz": 12, "val": "single"},
                            bottom={"sz": 12, "val": "single"},
                            start={"sz": 12, "val": "single"},
                            end={"sz": 12, "val": "single"})

            # Set font size for cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            
            # Set topic bold
            if cell == row_cells[1]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    # Add References section
    if "references" in json_data:
        doc.add_heading('References', level=2)
        i = 1
        for reference in json_data['references']:
            # Add each reference with its link as plain text
            ref_paragraph = doc.add_paragraph(style='BodyText')
            ref_paragraph.add_run(reference['reference']).italic = True
            if 'link' in reference:
                if reference['link'] != '':
                    # ref_paragraph.add_run("\nURL: ")
                    # ref_paragraph.add_run(reference['link'])
                    add_hyperlink(ref_paragraph, f'[{i}]', reference['link'])
                    i+=1
                    

    # Allow for titles containing .docx to still be saved as a word document ending in .docx instead of .docx.docx
    title = title.replace(".docx", "")
    doc.save(title+'.docx')
    return title+'.docx'

